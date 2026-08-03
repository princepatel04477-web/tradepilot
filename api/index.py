import os
import shutil
from pathlib import Path
from typing import Optional, List
from fastapi import FastAPI, File, UploadFile, Form, HTTPException, BackgroundTasks
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles

from app.core.constants import EXPORTS_DIR
from app.database import init_db, Repository
from app.models.contact import Contact
from app.models.template import EmailTemplate
from app.models.campaign import Campaign
from app.models.account import GmailAccount
from app.services.contact_service import ContactService
from app.services.template_service import template_service
from app.services.email_exporter import PerEmailExporter
from app.services.export_service import ExportService
from app.gmail.auth import GmailOAuthManager
from app.logger import logger

app = FastAPI(
    title="TradePilot - AI Outreach Platform",
    version="1.0.0",
    description="Remote Web Interface and API for TradePilot Outreach Engine"
)

# Safe database initialization
try:
    init_db()
except Exception as e:
    logger.warning(f"Database init warning: {e}")

BASE_DIR = Path(__file__).resolve().parent.parent

INDEX_HTML_CONTENT = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>TradePilot Remote — AI Email Outreach Platform</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap" rel="stylesheet">
    <style>
        :root {
            --bg-dark: #0f0f17;
            --card-bg: rgba(24, 24, 37, 0.85);
            --border-color: rgba(255, 255, 255, 0.1);
            --accent-blue: #89b4fa;
            --accent-green: #a6e3a1;
            --text-main: #cdd6f4;
            --text-muted: #a6adc8;
        }
        * { box-sizing: border-box; margin: 0; padding: 0; font-family: 'Inter', system-ui, sans-serif; }
        body { background-color: var(--bg-dark); color: var(--text-main); display: flex; height: 100vh; overflow: hidden; }
        .app-container { display: flex; width: 100vw; height: 100vh; }
        .sidebar { width: 250px; background: rgba(17, 17, 27, 0.95); border-right: 1px solid var(--border-color); padding: 24px 16px; display: flex; flex-direction: column; gap: 30px; }
        .brand h2 { color: var(--accent-blue); font-size: 20px; }
        .badge { display: inline-block; background: rgba(137, 180, 250, 0.15); color: var(--accent-blue); font-size: 11px; padding: 2px 8px; border-radius: 12px; margin-top: 4px; }
        .nav-menu { display: flex; flex-direction: column; gap: 8px; }
        .nav-btn { background: transparent; border: none; color: var(--text-muted); padding: 12px 16px; border-radius: 8px; text-align: left; font-size: 14px; cursor: pointer; transition: all 0.2s ease; }
        .nav-btn:hover { background: rgba(255, 255, 255, 0.05); color: var(--text-main); }
        .nav-btn.active { background: var(--accent-blue); color: #11111b; font-weight: 600; }
        .main-content { flex: 1; padding: 30px; overflow-y: auto; }
        .tab-content { display: none; flex-direction: column; gap: 20px; }
        .tab-content.active { display: flex; }
        .page-header h1 { font-size: 24px; font-weight: 700; }
        .subtitle { color: var(--text-muted); font-size: 13px; margin-top: 4px; }
        .stats-grid { display: grid; grid-template-columns: repeat(4, 1fr); gap: 16px; }
        .stat-card { background: var(--card-bg); border: 1px solid var(--border-color); padding: 20px; border-radius: 12px; backdrop-filter: blur(10px); }
        .card-title { color: var(--text-muted); font-size: 12px; text-transform: uppercase; }
        .stat-card h3 { font-size: 28px; color: var(--accent-blue); margin-top: 8px; }
        .card-panel { background: var(--card-bg); border: 1px solid var(--border-color); padding: 24px; border-radius: 12px; }
        .data-table { width: 100%; border-collapse: collapse; margin-top: 14px; }
        .data-table th, .data-table td { padding: 12px; text-align: left; border-bottom: 1px solid var(--border-color); }
        .data-table th { color: var(--text-muted); font-size: 13px; }
        .btn { padding: 10px 18px; border-radius: 8px; border: none; font-weight: 600; cursor: pointer; transition: background 0.2s; text-decoration: none; display: inline-block; }
        .btn-primary { background: var(--accent-blue); color: #11111b; }
        .btn-primary:hover { background: #b4befe; }
        .btn-secondary { background: rgba(255,255,255,0.1); color: var(--text-main); }
        .btn-success { background: var(--accent-green); color: #11111b; }
        .form-group { display: flex; flex-direction: column; gap: 6px; margin-bottom: 14px; }
        .form-group input, .form-group select, .form-group textarea { background: rgba(17, 17, 27, 0.8); border: 1px solid var(--border-color); color: var(--text-main); padding: 10px; border-radius: 6px; }
        .grid-2col { display: grid; grid-template-columns: 1fr 1fr; gap: 20px; }
        .form-row { display: flex; gap: 10px; }
        .checkbox-group { flex-direction: row; align-items: center; gap: 8px; }
        .modal { display:none; position:fixed; top:0; left:0; width:100vw; height:100vh; background:rgba(0,0,0,0.7); align-items:center; justify-content:center; z-index:100; }
        .modal-content { background:#181825; border:1px solid var(--border-color); padding:24px; border-radius:12px; width:450px; max-width:90%; }
        .all-in-one-panel { background: rgba(166, 227, 161, 0.05); border: 2px solid var(--accent-green); padding: 24px; border-radius: 12px; margin-bottom: 20px; }
    </style>
</head>
<body>
    <div class="app-container">
        <!-- Navigation Sidebar -->
        <aside class="sidebar">
            <div class="brand">
                <h2>✈ TradePilot</h2>
                <span class="badge">v1.3 All-In-One Form</span>
            </div>
            <nav class="nav-menu">
                <button class="nav-btn active" onclick="showTab('dashboard', event)">📊 Dashboard</button>
                <button class="nav-btn" onclick="showTab('campaigns', event)">🚀 All-In-One Campaign Sender</button>
                <button class="nav-btn" onclick="showTab('contacts', event)">👥 Contacts</button>
                <button class="nav-btn" onclick="showTab('gmail', event)">🔑 Gmail Accounts</button>
                <button class="nav-btn" onclick="showTab('templates', event)">📝 Templates</button>
                <button class="nav-btn" onclick="showTab('exports', event)">📄 Email PDF/TXT</button>
            </nav>
        </aside>

        <!-- Main Content View -->
        <main class="main-content">
            <!-- DASHBOARD TAB -->
            <section id="tab-dashboard" class="tab-content active">
                <header class="page-header">
                    <h1>Dashboard Overview</h1>
                    <p class="subtitle">Real-time metrics and remote campaign controls</p>
                </header>

                <div class="stats-grid">
                    <div class="stat-card">
                        <span class="card-title">Total Contacts</span>
                        <h3 id="stat-contacts">0</h3>
                    </div>
                    <div class="stat-card">
                        <span class="card-title">Campaigns</span>
                        <h3 id="stat-campaigns">0</h3>
                    </div>
                    <div class="stat-card">
                        <span class="card-title">Sent Today</span>
                        <h3 id="stat-sent">0</h3>
                    </div>
                    <div class="stat-card">
                        <span class="card-title">Success Rate</span>
                        <h3 id="stat-rate">100%</h3>
                    </div>
                </div>

                <div class="card-panel">
                    <h3>Recent Email Activity</h3>
                    <table class="data-table">
                        <thead>
                            <tr>
                                <th>ID</th>
                                <th>Recipient</th>
                                <th>Status</th>
                                <th>Timestamp</th>
                            </tr>
                        </thead>
                        <tbody id="activity-tbody">
                            <tr><td colspan="4" class="text-center">Loading activity logs...</td></tr>
                        </tbody>
                    </table>
                </div>
            </section>

            <!-- GMAIL ACCOUNTS TAB -->
            <section id="tab-gmail" class="tab-content">
                <header class="page-header">
                    <h1>Gmail OAuth Accounts Manager</h1>
                    <p class="subtitle">Connect sender accounts via official Google OAuth 2.0 (varunyainternational@gmail.com)</p>
                </header>

                <div class="card-panel">
                    <div style="display:flex; justify-content:space-between; align-items:center; margin-bottom:15px;">
                        <h3>Connected Accounts</h3>
                        <input type="file" id="json-file-input" style="display:none;" onchange="uploadCredentialsJson(event)">
                        <button class="btn btn-primary" onclick="document.getElementById('json-file-input').click()">🔑 Upload credentials.json to Connect Account</button>
                    </div>

                    <table class="data-table">
                        <thead>
                            <tr>
                                <th>ID</th>
                                <th>Sender Email</th>
                                <th>Status</th>
                                <th>Sent Today</th>
                                <th>Connected At</th>
                            </tr>
                        </thead>
                        <tbody id="accounts-tbody">
                            <tr><td colspan="5" class="text-center">Loading connected accounts...</td></tr>
                        </tbody>
                    </table>
                </div>
            </section>

            <!-- CAMPAIGNS TAB -->
            <section id="tab-campaigns" class="tab-content">
                <header class="page-header">
                    <h1>⚡ All-In-One Bulk Campaign Sender</h1>
                    <p class="subtitle">Upload your Word (.docx) template + Upload Contacts + Hit Send all right here in one step!</p>
                </header>

                <div class="all-in-one-panel">
                    <h3 style="color:var(--accent-green); margin-bottom:15px;">🚀 Single-Form Bulk Outreach Launcher</h3>
                    <form id="all-in-one-form" onsubmit="handleAllInOneSend(event)">
                        <div class="grid-2col">
                            <div>
                                <div class="form-group">
                                    <label style="font-weight:700;">1. Campaign Name</label>
                                    <input type="text" id="aio-name" required value="Frozen Shrimp Outreach">
                                </div>
                                <div class="form-group">
                                    <label style="font-weight:700; color:var(--accent-blue);">2. Upload Word (.docx) Email Format / Template</label>
                                    <input type="file" id="aio-docx-file" accept=".docx" style="background:#11111b; padding:8px;">
                                    <span style="font-size:11px; color:var(--text-muted);">Optional: Leave blank to use saved 'Frozen Shrimp Supply' template.</span>
                                </div>
                                <div class="form-group">
                                    <label style="font-weight:700; color:var(--accent-green);">3. Upload Contacts List (Excel / CSV)</label>
                                    <input type="file" id="aio-contacts-file" accept=".csv,.xlsx,.xls" style="background:#11111b; padding:8px;">
                                    <span style="font-size:11px; color:var(--text-muted);">Optional: Leave blank to send to all existing contacts in database.</span>
                                </div>
                            </div>

                            <div>
                                <div class="form-group">
                                    <label style="font-weight:700;">4. Target Email Mode</label>
                                    <select id="aio-target-mode" onchange="toggleAioTargetMode()">
                                        <option value="ALL">📢 BULK SEND: Send to ALL Uploaded Contacts in List</option>
                                        <option value="SINGLE">🎯 SINGLE TARGET: Send to One Specific Target Email</option>
                                    </select>
                                </div>
                                <div class="form-group" id="aio-single-email-container" style="display:none;">
                                    <label>Specific Target Email Address</label>
                                    <input type="email" id="aio-single-email" placeholder="e.g. enquiries@dibellacoffee.com">
                                </div>
                                <div class="form-group">
                                    <label style="font-weight:700;">5. Sender Email Account</label>
                                    <select id="camp-account" required></select>
                                </div>
                                <div class="form-group checkbox-group" style="margin-top:20px;">
                                    <input type="checkbox" id="aio-dryrun">
                                    <label for="aio-dryrun" style="font-weight:600;">Dry Run Mode (Simulate send)</label>
                                </div>
                            </div>
                        </div>

                        <button type="submit" class="btn btn-success" style="width:100%; font-size:18px; padding:16px; margin-top:15px; border-radius:10px; cursor:pointer;">
                            🚀 UPLOAD WORD FILE, IMPORT CONTACTS & BULK SEND NOW!
                        </button>
                    </form>
                </div>

                <div class="card-panel">
                    <h3>Sent Campaigns History</h3>
                    <table class="data-table">
                        <thead>
                            <tr>
                                <th>ID</th>
                                <th>Name</th>
                                <th>Status</th>
                                <th>Recipients</th>
                                <th>Action</th>
                            </tr>
                        </thead>
                        <tbody id="campaigns-tbody"></tbody>
                    </table>
                </div>
            </section>

            <!-- CONTACTS TAB -->
            <section id="tab-contacts" class="tab-content">
                <header class="page-header">
                    <h1>Contact Management</h1>
                </header>
                <div class="card-panel">
                    <div style="display:flex; gap:10px; margin-bottom:15px;">
                        <input type="file" id="contact-file-input" style="display:none;" onchange="uploadContactFile(event)">
                        <button class="btn btn-primary" onclick="document.getElementById('contact-file-input').click()">📁 Upload Excel / CSV (Bulk Contacts)</button>
                        <button class="btn btn-secondary" onclick="openAddContactModal()">➕ Add Custom Contact Email</button>
                    </div>
                    <table class="data-table">
                        <thead>
                            <tr>
                                <th>ID</th>
                                <th>Company</th>
                                <th>Name</th>
                                <th>Email</th>
                                <th>Country</th>
                            </tr>
                        </thead>
                        <tbody id="contacts-tbody"></tbody>
                    </table>
                </div>
            </section>

            <!-- TEMPLATES TAB -->
            <section id="tab-templates" class="tab-content">
                <header class="page-header">
                    <h1>Email Template & Word (.docx) Upload Center</h1>
                </header>
                <div class="card-panel">
                    <div style="background:rgba(166,227,161,0.1); border:1px solid var(--accent-green); padding:16px; border-radius:8px; margin-bottom:20px; display:flex; justify-content:space-between; align-items:center;">
                        <div>
                            <strong>📄 Upload Word Document (.docx) Template</strong>
                            <p style="font-size:12px; color:var(--text-muted);">Select any .docx file to automatically parse linebreaks, bullets, and signatures.</p>
                        </div>
                        <input type="file" id="docx-file-input" accept=".docx" style="display:none;" onchange="uploadDocxTemplate(event)">
                        <button class="btn btn-success" onclick="document.getElementById('docx-file-input').click()">📄 Upload Word (.docx) File</button>
                    </div>
                    <form id="template-form" onsubmit="handleCreateTemplate(event)">
                        <div class="form-group">
                            <label>Template Name</label>
                            <input type="text" id="tpl-name" required placeholder="Frozen Shrimp Supply">
                        </div>
                        <div class="form-group">
                            <label>Subject Line</label>
                            <input type="text" id="tpl-subject" required placeholder="Frozen Shrimp Supply">
                        </div>
                        <div class="form-group">
                            <label>Template Body (Jinja2 / Word Text)</label>
                            <textarea id="tpl-body" rows="10" required placeholder="Dear {{Contact}}, ..."></textarea>
                        </div>
                        <button type="submit" class="btn btn-primary">💾 Save Template</button>
                    </form>
                </div>
            </section>

            <!-- PER-EMAIL EXPORTS TAB -->
            <section id="tab-exports" class="tab-content">
                <header class="page-header">
                    <h1>Per-Email PDF & TXT Document Center</h1>
                    <p class="subtitle">Download individual formatted .pdf and .txt email copies one by one or as a ZIP bundle</p>
                </header>
                <div class="card-panel">
                    <h3>Generated Document Bundles</h3>
                    <p style="margin-bottom:15px; color:var(--text-muted);">Select a campaign to download all individual rendered email documents (.pdf and .txt):</p>
                    <div id="exports-list"></div>
                </div>
            </section>
        </main>
    </div>

    <!-- Add Contact Modal -->
    <div id="add-contact-modal" class="modal">
        <div class="modal-content">
            <h3 style="margin-bottom:15px;">➕ Add Custom Recipient Contact</h3>
            <form onsubmit="handleAddSingleContact(event)">
                <div class="form-group">
                    <label>Target Email Address *</label>
                    <input type="email" id="single-email" required placeholder="buyer@anycompany.com">
                </div>
                <div class="form-group">
                    <label>Contact Name</label>
                    <input type="text" id="single-name" placeholder="Procurement Manager">
                </div>
                <div class="form-group">
                    <label>Company Name</label>
                    <input type="text" id="single-company" placeholder="Global Imports LLC">
                </div>
                <div class="form-group">
                    <label>Country</label>
                    <input type="text" id="single-country" placeholder="United States">
                </div>
                <div style="display:flex; justify-content:flex-end; gap:10px; margin-top:15px;">
                    <button type="button" class="btn btn-secondary" onclick="closeAddContactModal()">Cancel</button>
                    <button type="submit" class="btn btn-primary">Save Contact</button>
                </div>
            </form>
        </div>
    </div>

    <script>
        document.addEventListener("DOMContentLoaded", () => {
            loadStats();
            loadActivityLogs();
            loadCampaigns();
            loadContacts();
            loadTemplates();
            loadAccounts();
        });

        function showTab(tabName, event) {
            document.querySelectorAll('.tab-content').forEach(el => el.classList.remove('active'));
            document.querySelectorAll('.nav-btn').forEach(el => el.classList.remove('active'));
            
            document.getElementById(`tab-${tabName}`).classList.add('active');
            if (event) event.target.classList.add('active');

            if (tabName === 'dashboard') loadStats();
            if (tabName === 'campaigns') loadCampaigns();
            if (tabName === 'contacts') loadContacts();
            if (tabName === 'gmail') loadAccounts();
            if (tabName === 'templates') loadTemplates();
            if (tabName === 'exports') loadExports();
        }

        function toggleAioTargetMode() {
            const mode = document.getElementById('aio-target-mode').value;
            document.getElementById('aio-single-email-container').style.display = (mode === 'SINGLE') ? 'block' : 'none';
        }

        async function loadStats() {
            try {
                const res = await fetch('/api/stats');
                const data = await res.json();
                document.getElementById('stat-contacts').innerText = data.total_contacts || 0;
                document.getElementById('stat-campaigns').innerText = data.total_campaigns || 0;
                document.getElementById('stat-sent').innerText = data.sent_today || 0;
                document.getElementById('stat-rate').innerText = `${data.success_rate || 100}%`;
            } catch (e) {
                console.error("Stats load failed", e);
            }
        }

        async function loadActivityLogs() {
            try {
                const res = await fetch('/api/logs?limit=10');
                const logs = await res.json();
                const tbody = document.getElementById('activity-tbody');
                if (!logs || logs.length === 0) {
                    tbody.innerHTML = '<tr><td colspan="4" class="text-center">No recent activity</td></tr>';
                    return;
                }
                tbody.innerHTML = logs.map(l => `
                    <tr>
                        <td>#${l.id}</td>
                        <td>${l.email}</td>
                        <td><span class="badge">${l.status}</span></td>
                        <td>${l.timestamp}</td>
                    </tr>
                `).join('');
            } catch (e) {
                console.error("Logs load failed", e);
            }
        }

        async function loadAccounts() {
            try {
                const res = await fetch('/api/accounts');
                const accounts = await res.json();
                const tbody = document.getElementById('accounts-tbody');
                const select = document.getElementById('camp-account');
                
                if (!accounts || accounts.length === 0) {
                    tbody.innerHTML = '<tr><td colspan="5" class="text-center">No connected Gmail accounts</td></tr>';
                    select.innerHTML = '<option value="">No accounts connected (Dry Run Only)</option>';
                    return;
                }
                
                tbody.innerHTML = accounts.map(a => `
                    <tr>
                        <td>#${a.id}</td>
                        <td><strong>${a.email}</strong></td>
                        <td><span class="badge" style="background:rgba(166,227,161,0.2); color:#a6e3a1;">${a.is_active ? 'Active OAuth' : 'Pending'}</span></td>
                        <td>${a.sent_today_count}</td>
                        <td>${a.created_at}</td>
                    </tr>
                `).join('');

                select.innerHTML = accounts.map(a => `<option value="${a.id}">${a.email} (ID #${a.id})</option>`).join('');
            } catch (e) {
                console.error("Accounts load failed", e);
            }
        }

        async function loadCampaigns() {
            try {
                const res = await fetch('/api/campaigns');
                const campaigns = await res.json();
                const tbody = document.getElementById('campaigns-tbody');
                if (!campaigns || campaigns.length === 0) {
                    tbody.innerHTML = '<tr><td colspan="5" class="text-center">No campaigns created yet</td></tr>';
                    return;
                }
                tbody.innerHTML = campaigns.map(c => `
                    <tr>
                        <td>#${c.id}</td>
                        <td>${c.name}</td>
                        <td>${c.status}</td>
                        <td>${c.sent_count} / ${c.total_recipients}</td>
                        <td><a href="/api/campaigns/${c.id}/export-bundle" class="btn btn-primary" style="font-size:12px; padding:4px 8px;">📦 PDF & TXT Zip</a></td>
                    </tr>
                `).join('');
            } catch (e) {
                console.error("Campaigns load failed", e);
            }
        }

        async function loadContacts() {
            try {
                const res = await fetch('/api/contacts');
                const contacts = await res.json();
                const tbody = document.getElementById('contacts-tbody');
                if (!contacts || contacts.length === 0) {
                    tbody.innerHTML = '<tr><td colspan="5" class="text-center">No contacts in database</td></tr>';
                    return;
                }
                tbody.innerHTML = contacts.map(c => `
                    <tr>
                        <td>#${c.id}</td>
                        <td>${c.company}</td>
                        <td>${c.contact_name}</td>
                        <td>${c.email}</td>
                        <td>${c.country}</td>
                    </tr>
                `).join('');
            } catch (e) {
                console.error("Contacts load failed", e);
            }
        }

        async function loadTemplates() {
            try {
                const res = await fetch('/api/templates');
                const templates = await res.json();
                if (templates.length > 0) {
                    document.getElementById('tpl-name').value = templates[0].name;
                    document.getElementById('tpl-subject').value = templates[0].subject;
                    document.getElementById('tpl-body').value = templates[0].body_content;
                }
            } catch (e) {
                console.error("Templates load failed", e);
            }
        }

        async function loadExports() {
            try {
                const res = await fetch('/api/campaigns');
                const campaigns = await res.json();
                const container = document.getElementById('exports-list');
                if (!campaigns || campaigns.length === 0) {
                    container.innerHTML = '<p>No campaigns found to export.</p>';
                    return;
                }
                container.innerHTML = campaigns.map(c => `
                    <div style="background: rgba(255,255,255,0.03); border:1px solid var(--border-color); padding:16px; margin-bottom:12px; border-radius:8px; display:flex; justify-content:space-between; align-items:center;">
                        <div>
                            <strong>Campaign #${c.id}: ${c.name}</strong>
                            <p style="color:var(--text-muted); font-size:12px;">Total Emails Rendered: ${c.total_recipients}</p>
                        </div>
                        <a href="/api/campaigns/${c.id}/export-bundle" class="btn btn-primary">⬇ Download Per-Email PDF & TXT Bundle (.ZIP)</a>
                    </div>
                `).join('');
            } catch (e) {
                console.error("Exports load failed", e);
            }
        }

        function openAddContactModal() {
            document.getElementById('add-contact-modal').style.display = 'flex';
        }

        function closeAddContactModal() {
            document.getElementById('add-contact-modal').style.display = 'none';
        }

        async function handleAddSingleContact(e) {
            e.preventDefault();
            const formData = new FormData();
            formData.append('email', document.getElementById('single-email').value);
            formData.append('contact_name', document.getElementById('single-name').value);
            formData.append('company', document.getElementById('single-company').value);
            formData.append('country', document.getElementById('single-country').value);

            try {
                const res = await fetch('/api/contacts/add', { method: 'POST', body: formData });
                alert("Contact added successfully!");
                closeAddContactModal();
                loadContacts();
                loadStats();
            } catch (err) {
                alert("Failed to add contact: " + err.message);
            }
        }

        async function uploadCredentialsJson(e) {
            const file = e.target.files[0];
            if (!file) return;
            const formData = new FormData();
            formData.append('file', file);
            try {
                const res = await fetch('/api/accounts/connect', { method: 'POST', body: formData });
                const data = await res.json();
                if (res.ok && data.email) {
                    alert(`Connected Gmail Account: ${data.email}`);
                    loadAccounts();
                } else {
                    alert("Account connection error: " + (data.detail || "Invalid credentials.json format"));
                }
            } catch (err) {
                alert("Account connection failed: " + err.message);
            }
        }

        async function handleAllInOneSend(e) {
            e.preventDefault();
            const formData = new FormData();
            formData.append('name', document.getElementById('aio-name').value);
            
            const docxInput = document.getElementById('aio-docx-file');
            if (docxInput.files[0]) formData.append('docx_file', docxInput.files[0]);

            const contactsInput = document.getElementById('aio-contacts-file');
            if (contactsInput.files[0]) formData.append('contacts_file', contactsInput.files[0]);

            const mode = document.getElementById('aio-target-mode').value;
            if (mode === 'SINGLE') {
                formData.append('target_email', document.getElementById('aio-single-email').value);
            }

            formData.append('is_dry_run', document.getElementById('aio-dryrun').checked);

            try {
                const res = await fetch('/api/campaigns/create-and-send', { method: 'POST', body: formData });
                const data = await res.json();
                alert(`🚀 Campaign Launched! Total Recipients: ${data.total_recipients}. Download per-email PDF/TXT bundle in Exports tab.`);
                loadCampaigns();
                loadStats();
            } catch (err) {
                alert("Campaign launch failed: " + err.message);
            }
        }

        async function uploadDocxTemplate(e) {
            const file = e.target.files[0];
            if (!file) return;
            const formData = new FormData();
            formData.append('file', file);
            try {
                const res = await fetch('/api/templates/upload-docx', { method: 'POST', body: formData });
                const data = await res.json();
                alert(`Uploaded Word Document Template: ${data.name}`);
                loadTemplates();
            } catch (err) {
                alert("Word template upload failed: " + err.message);
            }
        }

        async function handleCreateTemplate(e) {
            e.preventDefault();
            const formData = new FormData();
            formData.append('name', document.getElementById('tpl-name').value);
            formData.append('subject', document.getElementById('tpl-subject').value);
            formData.append('body_content', document.getElementById('tpl-body').value);

            try {
                const res = await fetch('/api/templates', { method: 'POST', body: formData });
                alert("Template saved successfully!");
                loadTemplates();
            } catch (err) {
                alert("Failed to save template: " + err.message);
            }
        }

        async function uploadContactFile(e) {
            const file = e.target.files[0];
            if (!file) return;
            const formData = new FormData();
            formData.append('file', file);
            try {
                const res = await fetch('/api/contacts/upload', { method: 'POST', body: formData });
                const data = await res.json();
                alert(`Imported ${data.inserted} contacts successfully!`);
                loadContacts();
                loadStats();
            } catch (err) {
                alert("Upload failed: " + err.message);
            }
        }
    </script>
</body>
</html>
"""

@app.get("/", response_class=HTMLResponse)
def read_root():
    return INDEX_HTML_CONTENT

@app.get("/api/stats")
def get_stats():
    try:
        return Repository.get_dashboard_stats()
    except Exception as e:
        return {"total_contacts": 0, "total_campaigns": 0, "sent_today": 0, "queued": 0, "failed": 0, "success_rate": 100.0}

@app.get("/api/accounts")
def get_accounts():
    accounts = Repository.get_accounts()
    return [
        {
            "id": a.id, "email": a.email, "display_name": a.display_name,
            "is_active": a.is_active, "sent_today_count": a.sent_today_count,
            "created_at": a.created_at
        } for a in accounts
    ]

@app.post("/api/accounts/connect")
async def connect_account_json(file: UploadFile = File(...)):
    temp_dir = EXPORTS_DIR / "uploads"
    temp_dir.mkdir(parents=True, exist_ok=True)
    temp_path = temp_dir / file.filename

    with open(temp_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    res = GmailOAuthManager.start_oauth_flow(str(temp_path))
    if not res:
        raise HTTPException(status_code=400, detail="OAuth authorization failed.")

    acc = GmailAccount(
        email=res["email"],
        display_name=res["email"],
        refresh_token_encrypted=res["encrypted_token"],
        is_active=True
    )
    acc_id = Repository.add_account(acc)
    return {"id": acc_id, "email": res["email"], "status": "connected"}

@app.get("/api/contacts")
def get_contacts(search: str = "", status: str = ""):
    contacts = Repository.get_contacts(search=search, status=status)
    return [
        {
            "id": c.id, "company": c.company, "contact_name": c.contact_name,
            "email": c.email, "country": c.country, "city": c.city,
            "phone": c.phone, "tags": c.tags, "status": c.status
        } for c in contacts
    ]

@app.post("/api/contacts/add")
def add_single_contact(
    email: str = Form(...),
    company: Optional[str] = Form(""),
    contact_name: Optional[str] = Form(""),
    country: Optional[str] = Form(""),
    city: Optional[str] = Form(""),
    phone: Optional[str] = Form("")
):
    c = Contact(
        email=email, company=company, contact_name=contact_name,
        country=country, city=city, phone=phone
    )
    count = Repository.add_contacts_batch([c])
    return {"status": "added", "inserted": count}

@app.post("/api/contacts/upload")
async def upload_contacts(file: UploadFile = File(...)):
    temp_dir = EXPORTS_DIR / "uploads"
    temp_dir.mkdir(parents=True, exist_ok=True)
    temp_path = temp_dir / file.filename

    with open(temp_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    res = ContactService.import_contacts_to_db(str(temp_path))
    return res

@app.get("/api/templates")
def get_templates():
    templates = Repository.get_templates()
    return [
        {
            "id": t.id, "name": t.name, "subject": t.subject,
            "body_content": t.body_content, "is_html": t.is_html,
            "variables": t.variables
        } for t in templates
    ]

@app.post("/api/templates/upload-docx")
async def upload_docx_template(file: UploadFile = File(...)):
    temp_dir = EXPORTS_DIR / "uploads"
    temp_dir.mkdir(parents=True, exist_ok=True)
    temp_path = temp_dir / file.filename

    with open(temp_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    doc_text = ""
    try:
        import docx
        doc = docx.Document(str(temp_path))
        paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        doc_text = "\n\n".join(paras)
    except Exception as e:
        logger.error(f"Failed to parse docx file: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to parse Word file: {e}")

    tpl_name = Path(file.filename).stem
    template = EmailTemplate(
        name=f"Word: {tpl_name}",
        subject="Frozen Shrimp Supply",
        body_content=doc_text,
        is_html=True,
        variables=["Contact", "Company"]
    )
    t_id = Repository.add_template(template)
    return {"id": t_id, "name": f"Word: {tpl_name}", "body_content": doc_text}

@app.post("/api/templates")
def create_template(name: str = Form(...), subject: str = Form(...), body_content: str = Form(...), is_html: bool = Form(True)):
    valid, err = template_service.validate_template_syntax(body_content)
    if not valid:
        raise HTTPException(status_code=400, detail=f"Syntax Error: {err}")

    vars_list = template_service.extract_variables(body_content)
    template = EmailTemplate(
        name=name, subject=subject, body_content=body_content,
        is_html=is_html, variables=vars_list
    )
    t_id = Repository.add_template(template)
    return {"id": t_id, "status": "created"}

@app.post("/api/campaigns/create-and-send")
async def create_and_send_all_in_one(
    name: str = Form("Frozen Shrimp Outreach"),
    docx_file: Optional[UploadFile] = File(None),
    contacts_file: Optional[UploadFile] = File(None),
    target_email: Optional[str] = Form(None),
    min_delay: float = Form(30.0),
    max_delay: float = Form(60.0),
    is_dry_run: bool = Form(False)
):
    template_id = 1
    # 1. Parse Word File if uploaded
    if docx_file and docx_file.filename:
        temp_dir = EXPORTS_DIR / "uploads"
        temp_dir.mkdir(parents=True, exist_ok=True)
        docx_path = temp_dir / docx_file.filename
        with open(docx_path, "wb") as buffer:
            shutil.copyfileobj(docx_file.file, buffer)
        
        try:
            import docx
            doc = docx.Document(str(docx_path))
            paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            doc_text = "\n\n".join(paras)
            
            tpl = EmailTemplate(
                name=f"Word: {Path(docx_file.filename).stem}",
                subject="Frozen Shrimp Supply",
                body_content=doc_text,
                is_html=True
            )
            template_id = Repository.add_template(tpl)
        except Exception as e:
            logger.error(f"Docx parse error: {e}")

    # 2. Import Contacts File if uploaded
    if contacts_file and contacts_file.filename:
        temp_dir = EXPORTS_DIR / "uploads"
        temp_dir.mkdir(parents=True, exist_ok=True)
        contacts_path = temp_dir / contacts_file.filename
        with open(contacts_path, "wb") as buffer:
            shutil.copyfileobj(contacts_file.file, buffer)
        ContactService.import_contacts_to_db(str(contacts_path))

    # 3. Determine Contacts Target
    if target_email and target_email.strip():
        email_clean = target_email.strip()
        c = Contact(email=email_clean, company="Target Prospect", contact_name="Prospect")
        Repository.add_contacts_batch([c])
        contacts = Repository.get_contacts(search=email_clean)
        contact_ids = [contacts[0].id] if contacts else []
    else:
        contacts = Repository.get_contacts(status="Active")
        contact_ids = [c.id for c in contacts]

    if not contact_ids:
        raise HTTPException(status_code=400, detail="No contacts found to send emails.")

    account = Repository.get_accounts()
    account_id = account[0].id if account else None

    campaign = Campaign(
        name=name,
        account_id=account_id,
        template_id=template_id,
        status="QUEUED",
        min_delay_sec=min_delay,
        max_delay_sec=max_delay,
        is_dry_run=is_dry_run,
        subject_override="Frozen Shrimp Supply",
        total_recipients=len(contact_ids)
    )

    campaign_id = Repository.create_campaign(campaign, contact_ids)
    export_res = PerEmailExporter.export_campaign_emails(campaign_id)

    return {
        "campaign_id": campaign_id,
        "status": "created",
        "total_recipients": len(contact_ids),
        "pdf_txt_export": export_res
    }

@app.get("/api/campaigns")
def get_campaigns():
    campaigns = Repository.get_campaigns()
    return [
        {
            "id": c.id, "name": c.name, "status": c.status,
            "total_recipients": c.total_recipients, "sent_count": c.sent_count,
            "failed_count": c.failed_count, "is_dry_run": c.is_dry_run,
            "created_at": c.created_at
        } for c in campaigns
    ]

@app.get("/api/campaigns/{campaign_id}/export-bundle")
def download_export_bundle(campaign_id: int):
    export_res = PerEmailExporter.export_campaign_emails(campaign_id)
    zip_path = Path(export_res["zip_bundle"])
    if not zip_path.exists():
        raise HTTPException(status_code=404, detail="Export bundle ZIP file not found.")
    return FileResponse(
        path=str(zip_path),
        filename=f"campaign_{campaign_id}_pdf_txt_bundle.zip",
        media_type="application/zip"
    )

@app.get("/api/logs")
def get_logs(limit: int = 100):
    logs = Repository.get_email_logs(limit=limit)
    return [
        {
            "id": l.id, "email": l.recipient_email, "status": l.status,
            "level": l.log_level, "details": l.details, "timestamp": l.timestamp
        } for l in logs
    ]
